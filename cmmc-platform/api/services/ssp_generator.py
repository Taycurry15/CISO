"""
SSP (System Security Plan) Generator Service

Generates CMMC System Security Plans from assessment data using AI-powered
narrative generation and professional DOCX templates.
"""

import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
import asyncpg
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class SSPSection(str, Enum):
    """SSP document sections"""
    SYSTEM_IDENTIFICATION = "system_identification"
    SYSTEM_DESCRIPTION = "system_description"
    SYSTEM_ENVIRONMENT = "system_environment"
    SYSTEM_BOUNDARY = "system_boundary"
    CONTROL_IMPLEMENTATION = "control_implementation"
    SYSTEM_INTERCONNECTIONS = "system_interconnections"
    PERSONNEL = "personnel"
    PLAN_MAINTENANCE = "plan_maintenance"


@dataclass
class SystemInfo:
    """System identification information"""
    system_name: str
    system_id: str
    system_type: str
    system_owner: str
    system_owner_email: str
    authorization_date: Optional[datetime]
    cmmc_level: int
    organization_name: str
    organization_address: str
    organization_phone: str
    organization_email: str
    data_types: List[str]
    mission: str
    system_description: str


@dataclass
class ControlImplementation:
    """Control implementation details for SSP"""
    control_id: str
    control_title: str
    control_requirement: str
    implementation_status: str  # Implemented, Partially Implemented, Planned, Not Applicable
    implementation_description: str
    evidence_description: str
    provider_inheritance: Optional[Dict[str, str]]
    responsible_role: str
    review_date: Optional[datetime]


@dataclass
class SSPMetadata:
    """SSP document metadata"""
    version: str
    date: datetime
    prepared_by: str
    reviewed_by: Optional[str]
    approved_by: Optional[str]
    classification: str  # CUI, FOUO, etc.


class SSPGenerator:
    """
    SSP Generator Service

    Generates professional System Security Plans from assessment data
    with AI-powered control narratives and evidence integration.
    """

    def __init__(
        self,
        db_pool: asyncpg.Pool,
        ai_service: Optional[Any] = None,
        template_path: Optional[str] = None
    ):
        """
        Initialize SSP generator

        Args:
            db_pool: Database connection pool
            ai_service: AI analysis service for narrative generation
            template_path: Path to SSP template DOCX file
        """
        self.db_pool = db_pool
        self.ai_service = ai_service

        if template_path:
            self.template_path = Path(template_path)
        else:
            # Default template location
            self.template_path = Path(__file__).parent.parent.parent / "templates" / "ssp_template.docx"

    async def generate_ssp(
        self,
        assessment_id: str,
        system_info: SystemInfo,
        metadata: SSPMetadata,
        include_provider_inheritance: bool = True,
        generate_narratives: bool = True
    ) -> BytesIO:
        """
        Generate complete SSP document

        Args:
            assessment_id: Assessment ID
            system_info: System identification info
            metadata: SSP metadata
            include_provider_inheritance: Include provider inheritance details
            generate_narratives: Use AI to generate control narratives

        Returns:
            BytesIO: Generated DOCX document
        """
        logger.info(f"Generating SSP for assessment {assessment_id}")

        # Create new document
        doc = Document()

        # Configure document formatting
        self._configure_document_styles(doc)

        # Add cover page
        self._add_cover_page(doc, system_info, metadata)
        doc.add_page_break()

        # Add table of contents placeholder
        self._add_table_of_contents(doc)
        doc.add_page_break()

        # 1. System Identification
        self._add_system_identification(doc, system_info, metadata)
        doc.add_page_break()

        # 2. System Description
        self._add_system_description(doc, system_info)
        doc.add_page_break()

        # 3. System Environment
        await self._add_system_environment(doc, assessment_id)
        doc.add_page_break()

        # 4. Control Implementation (main section)
        await self._add_control_implementation(
            doc,
            assessment_id,
            include_provider_inheritance,
            generate_narratives
        )
        doc.add_page_break()

        # 5. System Interconnections
        await self._add_system_interconnections(doc, assessment_id)
        doc.add_page_break()

        # 6. Personnel Roles
        self._add_personnel_roles(doc, system_info)
        doc.add_page_break()

        # 7. Plan Maintenance
        self._add_plan_maintenance(doc, metadata)

        # Save to BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        logger.info("SSP generation complete")
        return doc_bytes

    def _configure_document_styles(self, doc: Document):
        """Configure document styles"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading styles
        heading1 = doc.styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)

        heading2 = doc.styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 51, 102)

    def _add_cover_page(self, doc: Document, system_info: SystemInfo, metadata: SSPMetadata):
        """Add SSP cover page"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run("SYSTEM SECURITY PLAN")
        title_run.font.size = Pt(24)
        title_run.font.bold = True

        doc.add_paragraph()  # Spacing

        # System name
        system_name = doc.add_paragraph()
        system_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        system_name_run = system_name.add_run(system_info.system_name)
        system_name_run.font.size = Pt(18)
        system_name_run.font.bold = True

        doc.add_paragraph()  # Spacing

        # CMMC Level
        cmmc = doc.add_paragraph()
        cmmc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cmmc_run = cmmc.add_run(f"CMMC Level {system_info.cmmc_level}")
        cmmc_run.font.size = Pt(16)

        doc.add_paragraph()  # Spacing

        # Classification
        classification = doc.add_paragraph()
        classification.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        class_run = classification.add_run(metadata.classification)
        class_run.font.size = Pt(14)
        class_run.font.color.rgb = RGBColor(255, 0, 0)
        class_run.font.bold = True

        doc.add_paragraph()  # Spacing

        # Metadata table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = "Version:"
        table.rows[0].cells[1].text = metadata.version

        table.rows[1].cells[0].text = "Date:"
        table.rows[1].cells[1].text = metadata.date.strftime("%B %d, %Y")

        table.rows[2].cells[0].text = "Prepared By:"
        table.rows[2].cells[1].text = metadata.prepared_by

        table.rows[3].cells[0].text = "Reviewed By:"
        table.rows[3].cells[1].text = metadata.reviewed_by or "Pending"

        table.rows[4].cells[0].text = "Approved By:"
        table.rows[4].cells[1].text = metadata.approved_by or "Pending"

        table.rows[5].cells[0].text = "Organization:"
        table.rows[5].cells[1].text = system_info.organization_name

    def _add_table_of_contents(self, doc: Document):
        """Add table of contents placeholder"""
        doc.add_heading("Table of Contents", level=1)

        # Note: Word TOC requires manual update
        toc = doc.add_paragraph()
        toc.add_run("Right-click here and select 'Update Field' to generate table of contents.")
        toc.italic = True

    def _add_system_identification(self, doc: Document, system_info: SystemInfo, metadata: SSPMetadata):
        """Add system identification section"""
        doc.add_heading("1. System Identification", level=1)

        # System information table
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = "System Name:"
        table.rows[0].cells[1].text = system_info.system_name

        table.rows[1].cells[0].text = "System ID:"
        table.rows[1].cells[1].text = system_info.system_id

        table.rows[2].cells[0].text = "System Type:"
        table.rows[2].cells[1].text = system_info.system_type

        table.rows[3].cells[0].text = "CMMC Level:"
        table.rows[3].cells[1].text = f"Level {system_info.cmmc_level}"

        table.rows[4].cells[0].text = "System Owner:"
        table.rows[4].cells[1].text = system_info.system_owner

        table.rows[5].cells[0].text = "Owner Email:"
        table.rows[5].cells[1].text = system_info.system_owner_email

        table.rows[6].cells[0].text = "Organization:"
        table.rows[6].cells[1].text = system_info.organization_name

        table.rows[7].cells[0].text = "Address:"
        table.rows[7].cells[1].text = system_info.organization_address

        table.rows[8].cells[0].text = "Phone:"
        table.rows[8].cells[1].text = system_info.organization_phone

        table.rows[9].cells[0].text = "Email:"
        table.rows[9].cells[1].text = system_info.organization_email

    def _add_system_description(self, doc: Document, system_info: SystemInfo):
        """Add system description section"""
        doc.add_heading("2. System Description", level=1)

        doc.add_heading("2.1 Mission", level=2)
        doc.add_paragraph(system_info.mission)

        doc.add_heading("2.2 System Description", level=2)
        doc.add_paragraph(system_info.system_description)

        doc.add_heading("2.3 Data Types", level=2)
        for data_type in system_info.data_types:
            doc.add_paragraph(data_type, style='List Bullet')

    async def _add_system_environment(self, doc: Document, assessment_id: str):
        """Add system environment section"""
        doc.add_heading("3. System Environment", level=1)

        async with self.db_pool.acquire() as conn:
            # Get assessment scope
            scope = await conn.fetchrow("""
                SELECT scope, start_date, end_date
                FROM assessments
                WHERE id = $1
            """, assessment_id)

            if scope:
                doc.add_heading("3.1 Assessment Scope", level=2)
                doc.add_paragraph(scope['scope'] or "Not specified")

                doc.add_heading("3.2 Assessment Period", level=2)
                doc.add_paragraph(
                    f"Start Date: {scope['start_date'].strftime('%B %d, %Y') if scope['start_date'] else 'TBD'}\n"
                    f"End Date: {scope['end_date'].strftime('%B %d, %Y') if scope['end_date'] else 'TBD'}"
                )

    async def _add_control_implementation(
        self,
        doc: Document,
        assessment_id: str,
        include_provider_inheritance: bool,
        generate_narratives: bool
    ):
        """Add control implementation section (main SSP content)"""
        doc.add_heading("4. Control Implementation", level=1)

        async with self.db_pool.acquire() as conn:
            # Get all control findings for assessment
            findings = await conn.fetch("""
                SELECT
                    cf.control_id,
                    c.title as control_title,
                    c.requirement as control_requirement,
                    cd.name as domain_name,
                    cf.status,
                    cf.assessor_narrative,
                    cf.ai_confidence_score
                FROM control_findings cf
                JOIN controls c ON cf.control_id = c.id
                JOIN control_domains cd ON c.domain_id = cd.id
                WHERE cf.assessment_id = $1
                ORDER BY cd.name, cf.control_id
            """, assessment_id)

            current_domain = None

            for finding in findings:
                # Add domain header if new domain
                if finding['domain_name'] != current_domain:
                    current_domain = finding['domain_name']
                    doc.add_heading(f"4.{finding['domain_name']} - {current_domain}", level=2)

                # Control header
                doc.add_heading(
                    f"{finding['control_id']}: {finding['control_title']}",
                    level=3
                )

                # Status
                status_para = doc.add_paragraph()
                status_para.add_run("Implementation Status: ").bold = True
                status_para.add_run(self._map_status_to_implementation(finding['status']))

                # Control requirement
                doc.add_paragraph()
                req_para = doc.add_paragraph()
                req_para.add_run("Control Requirement: ").bold = True
                doc.add_paragraph(finding['control_requirement'], style='Quote')

                # Implementation description
                doc.add_paragraph()
                impl_para = doc.add_paragraph()
                impl_para.add_run("Implementation Description:").bold = True

                # Use AI narrative if available
                narrative = finding['assessor_narrative'] or "Implementation description not provided."
                doc.add_paragraph(narrative)

                # Provider inheritance (if applicable)
                if include_provider_inheritance:
                    inheritance = await self._get_provider_inheritance(conn, finding['control_id'])
                    if inheritance:
                        doc.add_paragraph()
                        provider_para = doc.add_paragraph()
                        provider_para.add_run("Provider Inheritance:").bold = True

                        for prov in inheritance:
                            doc.add_paragraph(
                                f"{prov['provider_name']} ({prov['responsibility']}): {prov['provider_narrative']}",
                                style='List Bullet'
                            )

                # Evidence references
                evidence = await self._get_evidence_references(conn, assessment_id, finding['control_id'])
                if evidence:
                    doc.add_paragraph()
                    ev_para = doc.add_paragraph()
                    ev_para.add_run("Evidence:").bold = True

                    for ev in evidence:
                        doc.add_paragraph(f"• {ev['title']}", style='List Bullet')

                doc.add_paragraph()  # Spacing between controls

    def _map_status_to_implementation(self, status: str) -> str:
        """Map assessment status to SSP implementation status"""
        mapping = {
            'Met': 'Implemented',
            'Partially Met': 'Partially Implemented',
            'Not Met': 'Planned',
            'Not Applicable': 'Not Applicable'
        }
        return mapping.get(status, 'Unknown')

    async def _get_provider_inheritance(
        self,
        conn: asyncpg.Connection,
        control_id: str
    ) -> List[Dict[str, Any]]:
        """Get provider inheritance for control"""
        rows = await conn.fetch("""
            SELECT
                po.provider_name,
                pci.responsibility,
                pci.provider_narrative
            FROM provider_control_inheritance pci
            JOIN provider_offerings po ON pci.provider_offering_id = po.id
            WHERE pci.control_id = $1
        """, control_id)

        return [dict(row) for row in rows]

    async def _get_evidence_references(
        self,
        conn: asyncpg.Connection,
        assessment_id: str,
        control_id: str
    ) -> List[Dict[str, str]]:
        """Get evidence references for control"""
        rows = await conn.fetch("""
            SELECT DISTINCT
                d.title,
                d.file_name
            FROM evidence e
            JOIN documents d ON e.document_id = d.id
            WHERE e.assessment_id = $1
              AND e.control_id = $2
        """, assessment_id, control_id)

        return [dict(row) for row in rows]

    async def _add_system_interconnections(self, doc: Document, assessment_id: str):
        """Add system interconnections section"""
        doc.add_heading("5. System Interconnections", level=1)

        doc.add_paragraph("This section documents connections to external systems and services.")

        # Placeholder - would query interconnections table if exists
        doc.add_paragraph("No external system interconnections documented.", style='List Bullet')

    def _add_personnel_roles(self, doc: Document, system_info: SystemInfo):
        """Add personnel roles section"""
        doc.add_heading("6. Personnel Roles and Responsibilities", level=1)

        doc.add_heading("6.1 System Owner", level=2)
        doc.add_paragraph(f"Name: {system_info.system_owner}")
        doc.add_paragraph(f"Email: {system_info.system_owner_email}")
        doc.add_paragraph("Responsibilities: Overall accountability for system security and compliance")

        doc.add_heading("6.2 Information System Security Officer (ISSO)", level=2)
        doc.add_paragraph("Responsibilities: Day-to-day security operations and monitoring")

        doc.add_heading("6.3 System Administrator", level=2)
        doc.add_paragraph("Responsibilities: System maintenance, patching, and configuration management")

    def _add_plan_maintenance(self, doc: Document, metadata: SSPMetadata):
        """Add plan maintenance section"""
        doc.add_heading("7. Plan Maintenance", level=1)

        doc.add_paragraph("This System Security Plan shall be reviewed and updated:")
        doc.add_paragraph("• Annually, or", style='List Bullet')
        doc.add_paragraph("• Whenever a significant change occurs to the system", style='List Bullet')
        doc.add_paragraph("• Prior to CMMC assessment", style='List Bullet')

        doc.add_paragraph()
        doc.add_paragraph(f"Current Version: {metadata.version}")
        doc.add_paragraph(f"Last Updated: {metadata.date.strftime('%B %d, %Y')}")


# Helper functions
async def generate_ssp_for_assessment(
    db_pool: asyncpg.Pool,
    assessment_id: str,
    system_info: SystemInfo,
    metadata: SSPMetadata,
    ai_service: Optional[Any] = None
) -> BytesIO:
    """
    Convenience function to generate SSP

    Args:
        db_pool: Database pool
        assessment_id: Assessment ID
        system_info: System information
        metadata: SSP metadata
        ai_service: Optional AI service

    Returns:
        BytesIO: Generated DOCX document
    """
    generator = SSPGenerator(db_pool, ai_service)
    return await generator.generate_ssp(assessment_id, system_info, metadata)
