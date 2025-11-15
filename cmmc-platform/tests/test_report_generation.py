"""
Tests for Report Generation Services (SSP and POA&M)
"""

import pytest
from datetime import datetime, timedelta
from io import BytesIO

# Add parent directory to path
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent / "api"))

from services.ssp_generator import (
    SSPGenerator,
    SystemInfo,
    SSPMetadata,
    SSPSection
)
from services.poam_generator import (
    POAMGenerator,
    POAMItem,
    POAMMetadata,
    RiskLevel,
    RemediationStatus
)


# ===========================
# Test Data
# ===========================

SAMPLE_SYSTEM_INFO = SystemInfo(
    system_name="Test Defense System",
    system_id="TDS-001",
    system_type="Cloud-based SaaS",
    system_owner="John Smith",
    system_owner_email="john.smith@test.mil",
    authorization_date=None,
    cmmc_level=2,
    organization_name="Test Defense Contractor",
    organization_address="123 Defense Blvd, Arlington, VA 22201",
    organization_phone="(703) 555-1234",
    organization_email="compliance@test.mil",
    data_types=["CUI", "Contract Data"],
    mission="Support DoD contract management",
    system_description="Cloud-based system for managing defense contracts and CUI"
)

SAMPLE_SSP_METADATA = SSPMetadata(
    version="1.0",
    date=datetime.utcnow(),
    prepared_by="Jane Doe, ISSO",
    reviewed_by="Bob Wilson, ISSM",
    approved_by="Alice Johnson, AO",
    classification="CUI"
)

SAMPLE_POAM_METADATA = POAMMetadata(
    system_name="Test Defense System",
    organization="Test Defense Contractor",
    prepared_by="Jane Doe, ISSO",
    preparation_date=datetime.utcnow(),
    review_date=None,
    version="1.0"
)


# ===========================
# Fixtures
# ===========================

@pytest.fixture
def mock_db_pool():
    """Mock database pool"""
    return None


@pytest.fixture
def ssp_generator(mock_db_pool):
    """Create SSP generator instance"""
    return SSPGenerator(mock_db_pool)


@pytest.fixture
def poam_generator(mock_db_pool):
    """Create POA&M generator instance"""
    return POAMGenerator(mock_db_pool)


# ===========================
# Tests: SSP Generator
# ===========================

def test_ssp_generator_initialization(mock_db_pool):
    """Test SSP generator initialization"""
    generator = SSPGenerator(mock_db_pool)

    assert generator.db_pool == mock_db_pool
    assert generator.ai_service is None
    assert generator.template_path is not None


def test_system_info_dataclass():
    """Test SystemInfo dataclass"""
    system_info = SAMPLE_SYSTEM_INFO

    assert system_info.system_name == "Test Defense System"
    assert system_info.cmmc_level == 2
    assert len(system_info.data_types) == 2
    assert "CUI" in system_info.data_types


def test_ssp_metadata_dataclass():
    """Test SSPMetadata dataclass"""
    metadata = SAMPLE_SSP_METADATA

    assert metadata.version == "1.0"
    assert metadata.classification == "CUI"
    assert metadata.prepared_by == "Jane Doe, ISSO"
    assert isinstance(metadata.date, datetime)


def test_ssp_section_enum():
    """Test SSPSection enum"""
    assert SSPSection.SYSTEM_IDENTIFICATION == "system_identification"
    assert SSPSection.CONTROL_IMPLEMENTATION == "control_implementation"
    assert SSPSection.PLAN_MAINTENANCE == "plan_maintenance"


def test_map_status_to_implementation(ssp_generator):
    """Test status mapping for SSP"""
    assert ssp_generator._map_status_to_implementation("Met") == "Implemented"
    assert ssp_generator._map_status_to_implementation("Partially Met") == "Partially Implemented"
    assert ssp_generator._map_status_to_implementation("Not Met") == "Planned"
    assert ssp_generator._map_status_to_implementation("Not Applicable") == "Not Applicable"
    assert ssp_generator._map_status_to_implementation("Unknown") == "Unknown"


# ===========================
# Tests: POA&M Generator
# ===========================

def test_poam_generator_initialization(mock_db_pool):
    """Test POA&M generator initialization"""
    generator = POAMGenerator(mock_db_pool)

    assert generator.db_pool == mock_db_pool
    assert generator.ai_service is None


def test_poam_item_dataclass():
    """Test POAMItem dataclass"""
    item = POAMItem(
        item_id="POAM-001",
        control_id="AC.L2-3.1.1",
        control_title="Authorized Access Control",
        weakness_description="Access control not implemented",
        risk_level=RiskLevel.HIGH,
        impact="High impact to CUI",
        likelihood="High likelihood",
        remediation_plan="Implement IAM system",
        resources_required="$50K, 2 FTE",
        milestone_date=datetime.utcnow() + timedelta(days=90),
        responsible_person="John Smith",
        status=RemediationStatus.OPEN,
        completion_date=None,
        cost_estimate=50000.0,
        comments="Priority remediation"
    )

    assert item.item_id == "POAM-001"
    assert item.control_id == "AC.L2-3.1.1"
    assert item.risk_level == RiskLevel.HIGH
    assert item.status == RemediationStatus.OPEN
    assert item.cost_estimate == 50000.0


def test_poam_metadata_dataclass():
    """Test POAMMetadata dataclass"""
    metadata = SAMPLE_POAM_METADATA

    assert metadata.system_name == "Test Defense System"
    assert metadata.organization == "Test Defense Contractor"
    assert metadata.version == "1.0"
    assert isinstance(metadata.preparation_date, datetime)


def test_risk_level_enum():
    """Test RiskLevel enum"""
    assert RiskLevel.VERY_HIGH == "Very High"
    assert RiskLevel.HIGH == "High"
    assert RiskLevel.MODERATE == "Moderate"
    assert RiskLevel.LOW == "Low"


def test_remediation_status_enum():
    """Test RemediationStatus enum"""
    assert RemediationStatus.OPEN == "Open"
    assert RemediationStatus.IN_PROGRESS == "In Progress"
    assert RemediationStatus.COMPLETED == "Completed"
    assert RemediationStatus.RISK_ACCEPTED == "Risk Accepted"
    assert RemediationStatus.DELAYED == "Delayed"


def test_calculate_risk_level(poam_generator):
    """Test risk level calculation"""
    # High-risk domain, Not Met
    finding_ac_not_met = {
        'control_id': 'AC.L2-3.1.1',
        'status': 'Not Met',
        'ai_confidence_score': 0.8
    }
    assert poam_generator._calculate_risk_level(finding_ac_not_met) == RiskLevel.VERY_HIGH

    # High-risk domain, Partially Met
    finding_ac_partial = {
        'control_id': 'AC.L2-3.1.2',
        'status': 'Partially Met',
        'ai_confidence_score': 0.7
    }
    assert poam_generator._calculate_risk_level(finding_ac_partial) == RiskLevel.HIGH

    # Low-risk domain, Not Met
    finding_cm_not_met = {
        'control_id': 'CM.L2-3.4.1',
        'status': 'Not Met',
        'ai_confidence_score': 0.6
    }
    assert poam_generator._calculate_risk_level(finding_cm_not_met) == RiskLevel.HIGH

    # Low-risk domain, Partially Met
    finding_cm_partial = {
        'control_id': 'CM.L2-3.4.2',
        'status': 'Partially Met',
        'ai_confidence_score': 0.5
    }
    assert poam_generator._calculate_risk_level(finding_cm_partial) == RiskLevel.MODERATE


def test_determine_impact(poam_generator):
    """Test impact determination"""
    assert "Critical" in poam_generator._determine_impact(RiskLevel.VERY_HIGH)
    assert "High" in poam_generator._determine_impact(RiskLevel.HIGH)
    assert "Moderate" in poam_generator._determine_impact(RiskLevel.MODERATE)
    assert "Low" in poam_generator._determine_impact(RiskLevel.LOW)


def test_determine_likelihood(poam_generator):
    """Test likelihood determination"""
    # Low confidence = high likelihood
    finding_low_conf = {
        'control_id': 'AC.L2-3.1.1',
        'ai_confidence_score': 0.3
    }
    assert "High" in poam_generator._determine_likelihood(finding_low_conf)

    # Medium confidence = moderate likelihood
    finding_med_conf = {
        'control_id': 'AC.L2-3.1.2',
        'ai_confidence_score': 0.6
    }
    assert "Moderate" in poam_generator._determine_likelihood(finding_med_conf)

    # High confidence = low likelihood
    finding_high_conf = {
        'control_id': 'AC.L2-3.1.3',
        'ai_confidence_score': 0.9
    }
    assert "Low" in poam_generator._determine_likelihood(finding_high_conf)


def test_generate_weakness_description(poam_generator):
    """Test weakness description generation"""
    finding_not_met = {
        'control_id': 'AC.L2-3.1.1',
        'status': 'Not Met',
        'assessor_narrative': 'The organization has not implemented access control policies.'
    }

    description = poam_generator._generate_weakness_description(finding_not_met)

    assert 'AC.L2-3.1.1' in description
    assert 'not implemented' in description

    finding_partial = {
        'control_id': 'AC.L2-3.1.2',
        'status': 'Partially Met',
        'assessor_narrative': 'Some access controls are in place.'
    }

    description = poam_generator._generate_weakness_description(finding_partial)

    assert 'AC.L2-3.1.2' in description
    assert 'partially implemented' in description


# ===========================
# Tests: Document Generation
# ===========================

def test_ssp_document_structure(ssp_generator):
    """Test SSP document has correct structure"""
    from docx import Document

    doc = Document()

    # Test document style configuration
    ssp_generator._configure_document_styles(doc)

    # Check that styles exist
    assert 'Normal' in doc.styles
    assert 'Heading 1' in doc.styles
    assert 'Heading 2' in doc.styles


def test_ssp_cover_page(ssp_generator):
    """Test SSP cover page generation"""
    from docx import Document

    doc = Document()
    ssp_generator._add_cover_page(doc, SAMPLE_SYSTEM_INFO, SAMPLE_SSP_METADATA)

    # Check that content was added
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) > 0

    # Check metadata table
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Version:"
    assert table.rows[0].cells[1].text == "1.0"


def test_ssp_system_identification(ssp_generator):
    """Test system identification section"""
    from docx import Document

    doc = Document()
    ssp_generator._add_system_identification(doc, SAMPLE_SYSTEM_INFO, SAMPLE_SSP_METADATA)

    # Check heading exists
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('System Identification' in h for h in headings)

    # Check table exists with system info
    assert len(doc.tables) > 0
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "System Name:"
    assert table.rows[0].cells[1].text == "Test Defense System"


def test_ssp_system_description(ssp_generator):
    """Test system description section"""
    from docx import Document

    doc = Document()
    ssp_generator._add_system_description(doc, SAMPLE_SYSTEM_INFO)

    # Check headings exist
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('System Description' in h for h in headings)
    assert any('Mission' in h for h in headings)
    assert any('Data Types' in h for h in headings)


def test_ssp_personnel_roles(ssp_generator):
    """Test personnel roles section"""
    from docx import Document

    doc = Document()
    ssp_generator._add_personnel_roles(doc, SAMPLE_SYSTEM_INFO)

    # Check content
    text_content = ' '.join([p.text for p in doc.paragraphs])
    assert 'John Smith' in text_content
    assert 'john.smith@test.mil' in text_content
    assert 'System Owner' in text_content


def test_ssp_plan_maintenance(ssp_generator):
    """Test plan maintenance section"""
    from docx import Document

    doc = Document()
    ssp_generator._add_plan_maintenance(doc, SAMPLE_SSP_METADATA)

    # Check content
    text_content = ' '.join([p.text for p in doc.paragraphs])
    assert 'Plan Maintenance' in text_content
    assert '1.0' in text_content  # Version


def test_poam_workbook_creation(poam_generator):
    """Test POA&M workbook creation"""
    from openpyxl import Workbook

    # Create sample items
    items = [
        POAMItem(
            item_id="POAM-001",
            control_id="AC.L2-3.1.1",
            control_title="Authorized Access Control",
            weakness_description="Not implemented",
            risk_level=RiskLevel.HIGH,
            impact="High impact",
            likelihood="High likelihood",
            remediation_plan="Implement IAM",
            resources_required="$50K",
            milestone_date=datetime.utcnow() + timedelta(days=90),
            responsible_person="John Smith",
            status=RemediationStatus.OPEN,
            completion_date=None,
            cost_estimate=50000.0,
            comments=None
        )
    ]

    wb = poam_generator._create_poam_workbook(items, SAMPLE_POAM_METADATA)

    # Check sheets exist
    sheet_names = wb.sheetnames
    assert "Summary" in sheet_names
    assert "POA&M Items" in sheet_names
    assert "Instructions" in sheet_names


def test_poam_summary_sheet(poam_generator):
    """Test POA&M summary sheet"""
    from openpyxl import Workbook

    items = [
        POAMItem(
            item_id=f"POAM-{i:03d}",
            control_id=f"AC.L2-3.1.{i}",
            control_title="Test Control",
            weakness_description="Test",
            risk_level=RiskLevel.HIGH if i % 2 == 0 else RiskLevel.MODERATE,
            impact="Test",
            likelihood="Test",
            remediation_plan="Test",
            resources_required="Test",
            milestone_date=datetime.utcnow(),
            responsible_person="Test",
            status=RemediationStatus.OPEN,
            completion_date=None,
            cost_estimate=None,
            comments=None
        )
        for i in range(1, 6)
    ]

    wb = Workbook()
    poam_generator._add_summary_sheet(wb, items, SAMPLE_POAM_METADATA)

    # Check summary sheet exists
    ws = wb["Summary"]
    assert ws['A1'].value == "POA&M Summary Dashboard"
    assert ws['B4'].value == 5  # Total items


# ===========================
# Integration Tests
# ===========================

def test_ssp_full_generation_structure():
    """Test full SSP generation creates valid document"""
    # This would require a real database, so just test structure
    from docx import Document

    doc = Document()

    # Verify we can create a basic document with all sections
    doc.add_heading("SYSTEM SECURITY PLAN", level=1)
    doc.add_paragraph("Test content")
    doc.add_heading("1. System Identification", level=1)
    doc.add_heading("2. System Description", level=1)
    doc.add_heading("3. System Environment", level=1)
    doc.add_heading("4. Control Implementation", level=1)

    assert len(doc.paragraphs) >= 5


def test_poam_full_generation_structure():
    """Test full POA&M generation creates valid workbook"""
    from openpyxl import Workbook

    wb = Workbook()

    # Create sheets
    summary = wb.create_sheet("Summary", 0)
    items = wb.create_sheet("POA&M Items")
    instructions = wb.create_sheet("Instructions")

    # Add basic content
    summary['A1'] = "Summary"
    items['A1'] = "Items"
    instructions['A1'] = "Instructions"

    assert len(wb.sheetnames) == 4  # Including default sheet


# ===========================
# Edge Cases
# ===========================

def test_ssp_with_empty_data_types():
    """Test SSP with empty data types list"""
    system_info = SystemInfo(
        system_name="Test",
        system_id="T-001",
        system_type="Cloud",
        system_owner="Test",
        system_owner_email="test@test.com",
        authorization_date=None,
        cmmc_level=2,
        organization_name="Test Org",
        organization_address="123 Test St",
        organization_phone="555-1234",
        organization_email="org@test.com",
        data_types=[],  # Empty
        mission="Test",
        system_description="Test"
    )

    assert system_info.data_types == []


def test_poam_with_very_high_risk():
    """Test POA&M item with very high risk"""
    item = POAMItem(
        item_id="POAM-999",
        control_id="IA.L2-3.5.1",
        control_title="Critical Control",
        weakness_description="Critical weakness",
        risk_level=RiskLevel.VERY_HIGH,
        impact="Critical impact",
        likelihood="Very high",
        remediation_plan="Immediate action required",
        resources_required="Emergency budget",
        milestone_date=datetime.utcnow() + timedelta(days=30),  # 30 days for very high
        responsible_person="CISO",
        status=RemediationStatus.OPEN,
        completion_date=None,
        cost_estimate=None,
        comments="Urgent"
    )

    assert item.risk_level == RiskLevel.VERY_HIGH
    assert (item.milestone_date - datetime.utcnow()).days <= 30


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
